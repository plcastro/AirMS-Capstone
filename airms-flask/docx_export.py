from datetime import datetime
from io import BytesIO

from docx import Document
from flask import Response


def build_docx_report(title, body):
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Generated: {datetime.utcnow().isoformat()}Z")
    doc.add_paragraph(str(body))
    buf = BytesIO()
    doc.save(buf)
    return Response(buf.getvalue(), mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
